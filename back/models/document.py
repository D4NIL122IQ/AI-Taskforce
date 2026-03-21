from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey, CheckConstraint
from sqlalchemy.orm import relationship
from datetime import datetime, timezone
from back.database import Base

TYPES_VALIDES = {"pdf", "txt", "docx", "md"}


class Document(Base):
    __tablename__ = "document"

    id_document  = Column(Integer,      primary_key=True, autoincrement=True)
    nom_fichier  = Column(String(255),  nullable=False)
    type_fichier = Column(String(10),   nullable=False)
    chemin       = Column(Text,         nullable=False)
    date_upload  = Column(DateTime,     default=lambda: datetime.now(timezone.utc))
    agent_id     = Column(Integer,      ForeignKey("agent.id_agent"), nullable=False)

    __table_args__ = (
        CheckConstraint(
            "type_fichier IN ('pdf','txt','docx','md')",
            name="ck_document_type_fichier"
        ),
    )

    agent = relationship("Agent", back_populates="documents")

    def __init__(self, nom_fichier: str, type_fichier: str, chemin: str,
                 agent_id: int, date_upload: datetime = None):
        if not nom_fichier:
            raise ValueError("Nom de fichier vide")
        if type_fichier not in TYPES_VALIDES:
            raise ValueError(f"Type de fichier non supporté : {type_fichier}")
        if not chemin:
            raise ValueError("Chemin vide")

        self.nom_fichier  = nom_fichier
        self.type_fichier = type_fichier
        self.chemin       = chemin
        self.agent_id     = agent_id
        self.date_upload  = date_upload or datetime.now(timezone.utc)

    def lire(self) -> str:
        """Extraction du contenu textuel selon le type de fichier."""
        import os
        if not os.path.exists(self.chemin):
            raise IOError(f"Fichier introuvable : {self.chemin}")

        if self.type_fichier == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(self.chemin)
            return "\n".join([page.extract_text() or "" for page in reader.pages])

        if self.type_fichier in {"txt", "md"}:
            with open(self.chemin, encoding="utf-8") as f:
                return f.read()

        if self.type_fichier == "docx":
            import docx
            doc = docx.Document(self.chemin)
            return "\n".join([p.text for p in doc.paragraphs])

        raise IOError(f"Type non supporté : {self.type_fichier}")

    def analyser(self) -> list:
        """Découpage en chunks pour l'indexation RAG (512 tokens, overlap 64)."""
        texte = self.lire()
        # Découpage simple par mots — à remplacer par RecursiveTextSplitter LangChain
        mots = texte.split()
        chunk_size, overlap = 512, 64
        chunks = []
        i = 0
        while i < len(mots):
            chunk_mots = mots[i:i + chunk_size]
            chunks.append({
                "text":        " ".join(chunk_mots),
                "source":      self.nom_fichier,
                "chunk_index": len(chunks)
            })
            i += chunk_size - overlap
        return chunks

    def toDict(self) -> dict:
        return {
            "id_document":  self.id_document,
            "nom_fichier":  self.nom_fichier,
            "type_fichier": self.type_fichier,
            "chemin":       self.chemin,
            "date_upload":  str(self.date_upload),
            "agent_id":     self.agent_id,
        }

    def __repr__(self):
        return f"<Document [{self.type_fichier}] {self.nom_fichier}>"