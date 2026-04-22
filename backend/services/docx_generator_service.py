"""
Service de génération de documents Word (.docx) à partir de contenu Markdown.

Utilisé par les agents ayant `generate_document=True` pour transformer
leur réponse textuelle en fichier Word structuré et téléchargeable.
"""

import os
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

# Dossier où sont stockés les documents générés
DOCS_DIR = Path("backend/generated_documents")
DOCS_DIR.mkdir(parents=True, exist_ok=True)


def generer_docx(contenu_markdown: str, titre: str = "Document généré", prefix: str = "doc") -> str:
    """
    Génère un fichier Word à partir d'un contenu Markdown.

    Args:
        contenu_markdown (str): Contenu brut en markdown (avec #, -, **, etc.)
        titre (str): Titre du document affiché en haut
        prefix (str): Préfixe du nom de fichier (ex: "rapport", "synthese")

    Returns:
        str: Nom du fichier généré (ex: "rapport_20260418_143022.docx")
    """
    doc = Document()

    # ── Titre principal ────────────────────────────────────────────
    heading = doc.add_heading(titre, level=0)
    heading.alignment = 1  # centré

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    date_run.italic = True
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    date_para.alignment = 1

    doc.add_paragraph()  # ligne vide

    # ── Parsing du markdown ────────────────────────────────────────
    lignes = contenu_markdown.strip().split("\n")

    for ligne in lignes:
        ligne_strip = ligne.strip()

        if not ligne_strip:
            continue

        # Titre niveau 1 (# Titre)
        if ligne_strip.startswith("# "):
            doc.add_heading(ligne_strip[2:].strip(), level=1)

        # Titre niveau 2 (## Titre)
        elif ligne_strip.startswith("## "):
            doc.add_heading(ligne_strip[3:].strip(), level=2)

        # Titre niveau 3 (### Titre)
        elif ligne_strip.startswith("### "):
            doc.add_heading(ligne_strip[4:].strip(), level=3)

        # Liste à puces (- ou *)
        elif ligne_strip.startswith("- ") or ligne_strip.startswith("* "):
            texte = ligne_strip[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _ajouter_texte_style(p, texte)

        # Liste numérotée (1. 2. etc.)
        elif re.match(r"^\d+\.\s", ligne_strip):
            texte = re.sub(r"^\d+\.\s", "", ligne_strip)
            p = doc.add_paragraph(style="List Number")
            _ajouter_texte_style(p, texte)

        # Paragraphe normal
        else:
            p = doc.add_paragraph()
            _ajouter_texte_style(p, ligne_strip)

    # ── Sauvegarde ─────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nom_fichier = f"{prefix}_{timestamp}.docx"
    chemin = DOCS_DIR / nom_fichier

    doc.save(chemin)
    print(f"[DOCX] Document généré : {chemin}")

    return nom_fichier


def _ajouter_texte_style(paragraphe, texte: str):
    """
    Ajoute du texte dans un paragraphe en gérant les styles Markdown inline :
    - **gras**
    - *italique*
    """
    # Pattern qui capture **...** ou *...* ou texte normal
    pattern = r"(\*\*[^\*]+\*\*|\*[^\*]+\*|[^\*]+)"
    morceaux = re.findall(pattern, texte)

    for morceau in morceaux:
        if morceau.startswith("**") and morceau.endswith("**"):
            run = paragraphe.add_run(morceau[2:-2])
            run.bold = True
        elif morceau.startswith("*") and morceau.endswith("*"):
            run = paragraphe.add_run(morceau[1:-1])
            run.italic = True
        else:
            paragraphe.add_run(morceau)


def get_document_path(nom_fichier: str) -> Path:
    """Retourne le chemin complet d'un document généré."""
    return DOCS_DIR / nom_fichier


def document_existe(nom_fichier: str) -> bool:
    """Vérifie qu'un document existe."""
    return (DOCS_DIR / nom_fichier).exists()