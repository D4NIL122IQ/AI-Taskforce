# backend/services/rag_service.py
"""
RAGService — Retrieval-Augmented Generation pour AI Taskforce.

PIPELINE COMPLET :
┌──────────────┐ → ┌──────────────┐ → ┌──────────────┐ → ┌──────────────┐ → ┌──────────────┐
│   Fichier    │   │ Extraction   │   │   Chunks     │   │ Embeddings   │   │  ChromaDB    │
│ (pdf/txt…)   │   │   texte      │   │ (chunk_size) │   │  (Ollama)    │   │ + metadata   │
└──────────────┘   └──────────────┘   └──────────────┘   └──────────────┘   └──────────────┘

  Lors d'une requête :
    Question utilisateur → Embedding (Ollama) → Retrieval (Similarity + MMR) → Top-K chunks → Post-traitement (filtrage + extraction + reranking) → LLM → Réponse

CONFIGURATION (.env) :
  TOKEN_PLEIADE=<votre_token>
  CHROMA_DIR=chroma_db       ← dossier de persistance ChromaDB (défaut)
  CHUNK_SIZE=500              ← taille des chunks en mots (défaut)
  CHUNK_OVERLAP=50            ← overlap entre chunks (défaut)
  EMBEDDING_MODEL="qwen3-embedding:4b"  ← modèle Pléiade pour les embeddings

PIPELINE COMPLET :
  Fichier → Chunks → Embeddings → Base vectorielle (Chroma)

Utilise une base Chroma persistante avec embeddings locaux.
Les documents sont filtrés par agent_id via les métadonnées.
"""

import os
from typing import List
from dotenv import load_dotenv
from pypdf import PdfReader

from backend.modeles.pleaide_embedding import PleiadesEmbeddings
from langchain_experimental.text_splitter import SemanticChunker
from langchain_core.documents import Document
from langchain_chroma import Chroma

load_dotenv()

# ── Configuration ──────────────────────────────────────────────────────────────

CHROMA_DIR = os.getenv("CHROMA_DIR", "chroma_db")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "500"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text-v2-moe:latest")


class RAGService:
    """
    Service RAG : indexation et recherche de documents via embeddings.

    Implémentation basée sur Chroma + LangChain.
    """
    

    def __init__(self):
        self._vectordb = None  

    # ──────────────────────────────────────────────────────────────────────────
    # PROPRIÉTÉ : base vectorielle (lazy)
    # ──────────────────────────────────────────────────────────────────────────
    @property
    def vectordb(self):
        if self._vectordb is None:
            embedding = PleiadesEmbeddings(model=EMBEDDING_MODEL)

            self._vectordb = Chroma(
                persist_directory=CHROMA_DIR,
                embedding_function=embedding
            )

        return self._vectordb

    # ──────────────────────────────────────────────────────────────────────────
    # ÉTAPE 1 : EXTRACTION DU TEXTE
    # ──────────────────────────────────────────────────────────────────────────
    

    def _decouper_en_chunks(self, documents: List[Document]) -> List[Document]:
        """Chunking parallèle — 1 thread par page/section."""
        splitter = SemanticChunker(
            embeddings=PleiadesEmbeddings(),
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=85,
        )

        def chunker_un_doc(doc: Document) -> List[Document]:
            return splitter.create_documents(
                [doc.page_content],
                metadatas=[doc.metadata]
            )

        chunks = []
        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(chunker_un_doc, doc) for doc in documents]
            for future in futures:
                chunks.extend(future.result())

        return chunks
    # ──────────────────────────────────────────────────────────────────────────
    # ÉTAPE 2 : Stratégie de découpage en chunks
    # ──────────────────────────────────────────────────────────────────────────

    def _extraire_en_documents_structures(self, chemin: str) -> List[Document]:
        """
        Extrait le texte EN CONSERVANT la structure comme métadonnées.
        Chaque section/paragraphe devient un Document LangChain avec contexte.
        """
        ext = os.path.splitext(chemin)[1].lower()

        if ext == ".pdf":
            return self._extraire_pdf_structure(chemin)
        elif ext == ".docx":
            return self._extraire_docx_structure(chemin)
        elif ext in (".txt", ".md"):
            return self._extraire_txt_structure(chemin)

    def _extraire_pdf_structure(self, chemin: str) -> List[Document]:
        """Extrait page par page avec numéro de page en métadonnée."""
        reader = PdfReader(chemin)
        docs = []
        for i, page in enumerate(reader.pages):
            texte = page.extract_text() or ""
            if texte.strip():
                docs.append(Document(
                    page_content=texte,
                    metadata={"page": i + 1, "source": chemin}
                ))
        return docs  # 1 Document par page → chunking ensuite

    def _extraire_docx_structure(self, chemin: str) -> List[Document]:
        """Regroupe par section Word (Heading → paragraphes suivants)."""
        from docx import Document as DocxDocument
        doc = DocxDocument(chemin)
        docs = []
        section_courante = ""
        titre_courant = "Introduction"

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Sauvegarde la section précédente
                if section_courante.strip():
                    docs.append(Document(
                        page_content=section_courante.strip(),
                        metadata={"titre_section": titre_courant, "source": chemin}
                    ))
                titre_courant = para.text
                section_courante = ""
            else:
                section_courante += para.text + "\n"

        # Dernière section
        if section_courante.strip():
            docs.append(Document(
                page_content=section_courante.strip(),
                metadata={"titre_section": titre_courant, "source": chemin}
            ))
        return docs

    def _extraire_txt_structure(self, chemin: str) -> List[Document]:
        """Découpe par double saut de ligne (paragraphes naturels)."""
        texte = open(chemin, "r", encoding="utf-8").read()
        paragraphes = [p.strip() for p in texte.split("\n\n") if p.strip()]
        return [
            Document(page_content=p, metadata={"source": chemin})
            for p in paragraphes
        ]

    # ──────────────────────────────────────────────────────────────────────────
    # ÉTAPE 3 : INDEXATION
    # ──────────────────────────────────────────────────────────────────────────
    def indexer_document(self, doc_id: int, agent_id: int, chemin_fichier: str) -> int:
        documents = self._extraire_en_documents_structures(chemin_fichier)
        chunks = self._decouper_en_chunks(documents)

        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "doc_id": int(doc_id),
                "agent_id": int(agent_id),
                "chunk_index": i,
                "chemin": chemin_fichier
            })

        # Stocker par lot
        stock_size = 50
        for i in range(0, len(chunks), stock_size):
            stock = chunks[i:i + stock_size]
            self.vectordb.add_documents(stock)
            print(f"[RAGService] Batch {i//stock_size + 1} stocké ({len(stock)} chunks)")

        print(f"[RAGService] Document {doc_id} indexé : {len(chunks)} chunks total")
        return len(chunks)

    # ──────────────────────────────────────────────────────────────────────────
    # RECHERCHE
    # ──────────────────────────────────────────────────────────────────────────

    def rechercher(self, agent_id: int, question: str, top_k: int = 5, use_post: bool = False) -> List[str]:

        if not question.strip():
            return []
        
        #parametrage en fonction de la quantité de donnés
        results = self.vectordb.get(
            where={"agent_id": int(agent_id)}
        )
        count_docs = len(results["ids"])

        fetch_k = min(30, count_docs)
        if count_docs > 30:
            fetch_k = count_docs // 3

        base_retriever = self.vectordb.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": top_k,
                "fetch_k": fetch_k,
                "lambda_mult": 0.5,
                "filter": {"agent_id": int(agent_id)}
            }
        )

        if use_post:
            docs = self.post_traitement(base_retriever, question)
        else:
            docs = base_retriever.invoke(question)

        return [doc.page_content for doc in docs]

    def post_traitement(self, retriever, question: str) -> list:
        docs = retriever.invoke(question)

        # Dédoublonnage
        unique_docs = []
        seen = set()
        for doc in docs:
            content = doc.page_content.strip()
            if content not in seen:
                seen.add(content)
                unique_docs.append(doc)

        # Reranking via Pléiade
        chunks = [doc.page_content for doc in unique_docs]
        chunks_rerankes = self._reranker_pleiade(question, chunks)

        # Reconstruire les docs dans le nouvel ordre
        chunk_to_doc = {doc.page_content: doc for doc in unique_docs}
        return [chunk_to_doc[c] for c in chunks_rerankes if c in chunk_to_doc]
    
    # ──────────────────────────────────────────────────────────────────────────
    # CONTEXTE POUR PROMPT
    # ──────────────────────────────────────────────────────────────────────────
    def contexte_pour_prompt(self, agent_id: int, question: str, top_k: int = 5) -> str:
            
        chunks = self.rechercher(agent_id, question, top_k)

        if not chunks:
            return ""

        lignes = ["=== Contexte documentaire pertinent ==="]
        for i, chunk in enumerate(chunks, 1):
            lignes.append(f"[Extrait {i}]\n{chunk}")
        lignes.append("=" * 40)

        return "\n\n".join(lignes)
    
    # on utilise LLM pour trier par ordre le retour de l'algorithme retrieval du RAG
    def _reranker_pleiade(self, question: str, chunks: list[str]) -> list[str]:
        liste = "\n\n".join(
            f"[{i+1}] {chunk}" for i, chunk in enumerate(chunks)
        )
        prompt = (
            f"Question : {question}\n\n"
            f"Voici {len(chunks)} extraits :\n{liste}\n\n"
            f"Classe ces extraits du plus pertinent au moins pertinent "
            f"pour répondre à la question. "
            f"Réponds uniquement avec les numéros dans l'ordre, séparés par des virgules. "
            f"Exemple : 3,1,4,2,5"
            f"Rien de plus, pas de commentaire"
        )
        from backend.modeles.requestLLM import chat
        reponse = chat(prompt, model="phi4:latest", max_tokens=20, temperature=0.0)
        
        # Parser l'ordre retourné
        try:
            ordre = [int(n.strip()) - 1 for n in reponse.split(",")]
            return [chunks[i] for i in ordre if 0 <= i < len(chunks)]
        except Exception:
            return chunks  # fallback : ordre original

    # ──────────────────────────────────────────────────────────────────────────
    # SUPPRESSION
    # ──────────────────────────────────────────────────────────────────────────
    def supprimer_document(self, doc_id: int):
        self.vectordb._collection.delete(where={"doc_id": doc_id})

    def supprimer_agent(self, agent_id: int):
        self.vectordb._collection.delete(where={"agent_id": agent_id})